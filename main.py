from flask import Flask, render_template, request, jsonify, send_from_directory
from docx import Document
import re
import os
import base64
import json
import hashlib
from lxml import etree
from pathlib import Path
import glob
import threading
import queue
import time
import traceback
import logging
import uuid
from werkzeug.utils import secure_filename

from dotenv import load_dotenv
load_dotenv()

from ai_engine import generate_response, check_connection, get_model_info, validate_and_fix_xml
from rag_engine import find_similar_example, add_example, get_db_stats, get_rag_db, save_rag_db, remove_text_tag_from_xml

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    handlers=[
        logging.FileHandler("app_audit.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("SecureApp")

try:
    from textual_grounding import pipeline_verify_and_fix as textual_verify_and_fix
    TEXTUAL_GROUNDING_AVAILABLE = True
except ImportError:
    TEXTUAL_GROUNDING_AVAILABLE = False
    logger.warning("textual_grounding.py не найден")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB

articles_db = {}
current_md_text = ""
processing_queue = queue.Queue()
processing_status = {"active": False, "current": 0, "total": 0, "results": {}, "template_id": None}

UPLOAD_FOLDER = "uploads"
MEMORY_FOLDER = "memory"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(MEMORY_FOLDER, exist_ok=True)

WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

DEFAULT_INSTRUCTION = """Ты — ассистент-библиограф. Твоя задача — извлечь метаданные из научной статьи в виде XML, строго следуя структуре примера. Тебе будет передан полный текст статьи (возможно, с остатками вёрстки, колонтитулами, ссылками).

Твоя задача:
1) выделить все необходимые поля (страницы, тип статьи, авторов с их данными, названия, аннотации, DOI, УДК, ключевые слова, список литературы и прочее, если таковые будут);
2) для каждого автора определить данные на русском и английском языках (если в тексте присутствуют оба варианта, иначе оставить только один lang), если таковые имеются;
3) сформировать на выходе XML-документ, в точности соответствующий приведённой ниже схеме и по стилю похожий на пример. Не добавляй никаких комментариев, только XML. Не оборачивай XML в кодовые блоки (не используй ```xml ... ```), только чистый XML-документ.

Запрещено:
1) какие-либо галлюцинации;
2) выдумывать, изменять, обобщать текст;
3) выдумывать, изменять тэги."""

TEXT_TAG_NAME = "text"

def validate_template_id(template_id):
    if not template_id or not isinstance(template_id, str):
        return False
    return bool(re.match(r'^[a-zA-Z0-9_-]+$', template_id))

def docx_to_markdown(filepath):
    doc = Document(filepath)
    md_lines = []
    for element in doc.element.body:
        if element.tag == f'{{{WORD_NS}}}p':
            p = next((p for p in doc.paragraphs if p._element is element), None)
            if not p: continue
            prefix = ""
            style_name = p.style.name if p.style else ""
            if any(h in style_name for h in ['Heading 1', 'Заголовок 1']): prefix = "# "
            elif any(h in style_name for h in ['Heading 2', 'Заголовок 2']): prefix = "## "
            elif any(h in style_name for h in ['Heading 3', 'Заголовок 3']): prefix = "### "
            elif 'List' in style_name: prefix = "- "

            paragraph_text = prefix
            has_content = False
            for child in element:
                text_found = "".join(t.text or "" for t in child.iter(f'{{{WORD_NS}}}t'))
                images_found = []
                for blip in child.iter(f'{{{DRAW_NS}}}blip'):
                    embed_id = blip.get(f'{{{REL_NS}}}embed')
                    if embed_id:
                        try:
                            image_part = doc.part.related_parts[embed_id]
                            img_base64 = base64.b64encode(image_part.blob).decode('utf-8')
                            ext = image_part.content_type.split('/')[-1]
                            if ext == 'jpeg': ext = 'jpg'
                            images_found.append(f"\n![image](data:image/{ext};base64,{img_base64})\n")
                        except: pass
                if text_found: has_content = True; paragraph_text += text_found
                if images_found: has_content = True; paragraph_text += "".join(images_found)
            if has_content or prefix == "- ": md_lines.append(paragraph_text + "\n")

        elif element.tag == f'{{{WORD_NS}}}tbl':
            table = next((t for t in doc.tables if t._element is element), None)
            if not table: continue
            for i, row in enumerate(table.rows):
                cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
                if i == 0: md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            md_lines.append("\n")
    return "\n".join(md_lines)

def get_base_instruction():
    filepath = os.path.join(MEMORY_FOLDER, "base_instruction.json")
    if not os.path.exists(filepath):
        return DEFAULT_INSTRUCTION
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data.get('instruction', DEFAULT_INSTRUCTION)
    except Exception:
        return DEFAULT_INSTRUCTION

def save_base_instruction(instruction):
    filepath = os.path.join(MEMORY_FOLDER, "base_instruction.json")
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump({'instruction': instruction}, f, ensure_ascii=False, indent=2)

DEFAULT_SPLIT_PATTERN = r'(Формат цитирования\s*:\s*.*?DOI\s*:\s*10\.\d{4,9}/[^\s]+)'

def split_journal(md_text, custom_pattern=None, periodicity=1, position='bottom'):
    """
    Разбивает журнал на статьи по паттерну с учётом периодичности.
    
    Args:
        md_text: текст журнала
        custom_pattern: regex-паттерн
        periodicity: сколько раз маркер встречается в одной статье (1, 2, 3...)
        position: где маркер ('top' = в начале статьи, 'bottom' = в конце, 'both' = пропускаем чётные)
    """
    pattern_str = custom_pattern if custom_pattern else DEFAULT_SPLIT_PATTERN
    if len(pattern_str) > 500:
        logger.warning(f"Отклонен слишком длинный regex-паттерн ({len(pattern_str)} симв.)")
        raise ValueError("Паттерн слишком сложный")
    
    try:
        citation_pattern = re.compile(pattern_str, re.DOTALL | re.IGNORECASE)
    except re.error as e:
        logger.warning(f"[Split] Неверный regex '{pattern_str}': {e}. Использую дефолтный.")
        citation_pattern = re.compile(DEFAULT_SPLIT_PATTERN, re.DOTALL | re.IGNORECASE)
    
    matches = list(citation_pattern.finditer(md_text))
    
    if not matches:
        print(f"[Split] Паттерн не найден в документе")
        return [{"title": "Статья 1", "content": md_text}]
    
    print(f"[Split] Найдено {len(matches)} вхождений паттерна, периодичность={periodicity}, позиция={position}")
    
    split_points = []
    
    if position == 'top':
        for i, match in enumerate(matches):
            if i % periodicity == 0:
                split_points.append(('before', match.start()))
    elif position == 'bottom':
        for i, match in enumerate(matches):
            if (i + 1) % periodicity == 0:  # каждое k-ное
                split_points.append(('after', match.end()))
    else:
        for i, match in enumerate(matches):
            if (i + 1) % periodicity == 0:
                split_points.append(('after', match.end()))
    
    split_points.sort(key=lambda x: x[1])
    
    articles = []
    prev_end = 0
    
    for i, (cut_type, cut_pos) in enumerate(split_points):
        if cut_pos > prev_end:
            article_content = md_text[prev_end:cut_pos].strip()
            if article_content:
                articles.append({
                    "title": f"Статья {len(articles) + 1}",
                    "content": article_content
                })
        prev_end = cut_pos
    
    if prev_end < len(md_text):
        last_content = md_text[prev_end:].strip()
        if last_content:
            articles.append({
                "title": f"Статья {len(articles) + 1}",
                "content": last_content
            })
    
    if not articles:
        return [{"title": "Статья 1", "content": md_text}]
    
    print(f"[Split] Разбито на {len(articles)} статей")
    return articles

def remove_dynamic_block(content, start_words, end_words):
    if not start_words or not end_words:
        return content

    content_lower = content.lower()
    earliest_start, start_len = -1, 0

    for sw in start_words:
        idx = content_lower.find(sw.lower())
        if idx != -1 and (earliest_start == -1 or idx < earliest_start):
            earliest_start, start_len = idx, len(sw)

    if earliest_start == -1:
        return content

    after_start = content[earliest_start + start_len:]
    earliest_end = -1

    for ew in end_words:
        esc = re.escape(ew)
        spaced = esc.replace(r'\s+', r'\\s+')
        pattern = f"{spaced}(?![а-яА-ЯёЁ,;:])"
        match = re.search(pattern, after_start, re.IGNORECASE)
        if match and (earliest_end == -1 or match.start() < earliest_end):
            earliest_end = match.start()

    if earliest_end == -1:
        return content

    before = content[:earliest_start].rstrip()
    after = after_start[earliest_end:].lstrip()
    return (before + "\n\n" + after).strip()

def inject_cut_text_into_xml(xml_text, cut_text):
    if not cut_text or not xml_text:
        return xml_text

    escaped = (cut_text
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
        .replace("'", '&apos;')
    )

    escaped = escaped.strip()

    tag = TEXT_TAG_NAME
    pattern = rf'(</references>\s*)(</article>)'
    replacement = f'\\1<{tag} lang="RUS">\n{escaped}\n</{tag}>\n\\2'
    result = re.sub(pattern, replacement, xml_text, flags=re.DOTALL)

    if result == xml_text and '</article>' in xml_text:
        result = xml_text.replace('</article>', f'<{tag} lang="RUS">\n{escaped}\n</{tag}>\n</article>')

    return result

def extract_text_tag_content(xml_text, tag_name="text"):
    if not xml_text or not tag_name:
        return xml_text, None
    
    pattern = rf'<{tag_name}[^>]*>(.*?)</{tag_name}>'
    match = re.search(pattern, xml_text, flags=re.DOTALL)
    
    if match:
        extracted = match.group(1)
        slim_xml = re.sub(pattern, '', xml_text, flags=re.DOTALL)
        slim_xml = re.sub(r'\n\s*\n\s*\n+', '\n\n', slim_xml).strip()
        return slim_xml, extracted
    
    return xml_text, None

def generate_few_shot_prompt(template, article_text, instruction=None, use_rag=False):
    instr = instruction or DEFAULT_INSTRUCTION
    rag_info = {"used": False, "similarity": 0.0, "example_id": None, "source": "none"}
    
    if use_rag:
        # === РЕЖИМ 1: АВТО (RAG) ===
        rag_example = find_similar_example(article_text, threshold=0.15)
        if rag_example:
            before = rag_example["before"]
            after = rag_example["after"]
            rag_info = {
                "used": True,
                "similarity": rag_example["similarity"],
                "example_id": rag_example["id"],
                "source": "rag"
            }
            print(f"[Prompt] 🤖 АВТО-РЕЖИМ: RAG нашёл пример (similarity={rag_example['similarity']:.2f})")
        else:
            best_candidate = find_similar_example(article_text, threshold=0.0)
            if best_candidate:
                preview = best_candidate["before"][:500].replace('\n', ' ')
                print(f"[Prompt] АВТО-РЕЖИМ: RAG не нашёл примеров выше порога 0.15")
                print(f"[Prompt] Лучший кандидат (similarity={best_candidate['similarity']:.3f}):")
                print(f"[Prompt] ID: {best_candidate['id'][:16]}...")
                print(f"[Prompt] Начало (500 симв.): {preview}...")
            else:
                print(f"[Prompt] ⚠️ АВТО-РЕЖИМ: RAG-база пуста.")
            
            print(f"[Prompt] Использую только базовую инструкцию.")
            rag_info = {"used": False, "similarity": 0.0, "example_id": None, "source": "instruction_only"}
            prompt = f"""{instr}\n\nТекст статьи для обработки:\n{article_text}"""
            return prompt, rag_info
    else:
        if not template or not template.get('before') or not template.get('after'):
            raise ValueError("Шаблон не найден или не содержит примеров (before/after)")
        
        before = template['before']
        after = template['after']
        rag_info = {
            "used": False,
            "similarity": 0.0,
            "example_id": None,
            "source": "template"
        }
        print(f"[Prompt] СТРОГИЙ РЕЖИМ: Использую шаблон '{template.get('name', 'unknown')}' (RAG отключён)")

    prompt = f"""{instr}

Пример:
Ниже приведён исходный текст статьи и правильный XML-результат, который должен получиться после его обработки.

Исходный текст статьи (пример):
{before}

XML-результат (образец):
{after}

Теперь обработай следующий текст статьи по аналогии с примером.

Текст статьи для обработки:
{article_text}"""
    
    return prompt, rag_info

def get_prompt_history(template_id):
    if not validate_template_id(template_id):
        logger.warning(f"Попытка Path Traversal в get_prompt_history: {template_id}")
        return {"versions": [], "active": 1}
    filepath = os.path.join(MEMORY_FOLDER, f"prompts_{template_id}.json")
    fallback = {
        "versions": [{"v": 1, "instruction": DEFAULT_INSTRUCTION, "timestamp": "Изначальная", "remark": "Базовая версия (V1)"}],
        "active": 1
    }
    if not os.path.exists(filepath):
        return fallback
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if 'versions' not in data or not data['versions']:
                raise ValueError("Empty history")
            return data
    except Exception as e:
        print(f"[Warning] Ошибка чтения истории промптов: {e}")
        return fallback


def save_prompt_history(template_id, history):
    if not validate_template_id(template_id):
        logger.error(f"Попытка записать историю для невалидного ID: {template_id}")
        return
    filepath = os.path.join(MEMORY_FOLDER, f"prompts_{template_id}.json")
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def restart_parsing(article_ids, template_id):
    for aid in article_ids:
        if aid in articles_db:
            articles_db[aid]["parsed"] = None
            if aid in processing_status["results"]:
                del processing_status["results"][aid]
            
            content = articles_db[aid].get("edited_content") or articles_db[aid]["filtered"]
            processing_queue.put((aid, content, template_id))
            print(f"[Restart] 🔄 {aid} поставлена в очередь (template={template_id})")
    
    processing_status["active"] = True
    processing_status["current"] = 0
    processing_status["total"] = len(article_ids)
        

def run_ai_inference(prompt):
    return generate_response(prompt)

def ensure_valid_xml(xml_str):
    if not xml_str or not xml_str.strip():
        return xml_str
    
    cleaned = re.sub(r'```xml\s*', '', xml_str)
    cleaned = re.sub(r'```\s*', '', cleaned).strip()
    cleaned = re.sub(r'&(?![a-zA-Z]+;|#[0-9]+;|#x[0-9a-fA-F]+;)', '&amp;', cleaned)
    
    if not cleaned:
        return xml_str
    
    try:
        from lxml import etree
        secure_parser = etree.XMLParser(resolve_entities=False, no_network=True)
        root = etree.fromstring(cleaned.encode('utf-8'), parser=secure_parser)
        return xml_str
    except Exception:
        pass

    try:
        from lxml import etree
        wrapped = f"<_DOC_>{cleaned}</_DOC_>"
        root = etree.fromstring(wrapped.encode('utf-8'), parser=secure_parser)
        return xml_str
    except Exception:
        pass
    
    try:
        from lxml import etree
        parser = etree.XMLParser(recover=True, encoding='utf-8', resolve_entities=False, no_network=True)
        wrapped = f"<_DOC_>{cleaned}</_DOC_>"
        root = etree.fromstring(wrapped.encode('utf-8'), parser)
        
        if root is None or len(root) == 0:
            return xml_str
        
        parts = []
        for child in root:
            parts.append(etree.tostring(child, encoding='unicode', pretty_print=True))
        fixed_xml = '\n'.join(parts).strip()
        
        print(f"[XML Fix] Валидатор исправил XML (было {len(xml_str)}, стало {len(fixed_xml)} симв.)")
        return fixed_xml
        
    except Exception as e:
        logger.error(f"[XML Fix] Не удалось автоисправить XML: {e}")
        return xml_str

def sanitize_final_xml(xml_str):
    if not xml_str or not xml_str.strip():
        return xml_str
    
    try:
        from lxml import etree
        
        cleaned = re.sub(r'```xml\s*', '', xml_str)
        cleaned = re.sub(r'```\s*', '', cleaned).strip()
        cleaned = re.sub(r'&(?![a-zA-Z]+;|#[0-9]+;|#x[0-9a-fA-F]+;)', '&amp;', cleaned)
        
        root = None
        is_wrapped = False
        
        try:
            secure_parser = etree.XMLParser(resolve_entities=False, no_network=True)
            root = etree.fromstring(cleaned.encode('utf-8'), parser=secure_parser)
        except etree.XMLSyntaxError:
            try:
                wrapped = f"<_SANITIZE_>{cleaned}</_SANITIZE_>"
                root = etree.fromstring(wrapped.encode('utf-8'), parser=secure_parser)
                is_wrapped = True
            except etree.XMLSyntaxError:
                parser = etree.XMLParser(recover=True, encoding='utf-8', resolve_entities=False, no_network=True)
                try:
                    wrapped = f"<_SANITIZE_>{cleaned}</_SANITIZE_>"
                    root = etree.fromstring(wrapped.encode('utf-8'), parser)
                    is_wrapped = True
                except:
                    return xml_str
        
        if root is None:
            return xml_str
        
        list_marker_pattern = re.compile(
            r'^\s*(?:'
            r'[-*•·▪▸►]\s*'     # -, *, •, ·, ▪, ▸, ►
            r'|\d+[.)]\s*'       # 1., 1), 12.
            r'|\(\d+\)\s*'       # (1), (12)
            r'|[a-zа-яё]\)\s*'   # а), б), a)
            r'|[ivxlcdm]+[.)]\s*' # i., ii., iii. (римские)
            r')+'
        )
        
        reference_tags = {'reference', 'ref', 'mixed-citation', 'element-citation', 
                         'citation', 'bibl', 'bib', 'source'}
        
        fixes_count = 0
        
        for el in root.iter():
            if el.text and el.text.strip():
                original_text = el.text
                tag_name = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                
                if tag_name.lower() in reference_tags:
                    cleaned_text = list_marker_pattern.sub('', el.text).strip()
                    if cleaned_text != original_text.strip():
                        el.text = cleaned_text
                        fixes_count += 1
        
        if is_wrapped:
            parts = []
            for child in root:
                child_str = etree.tostring(child, encoding='unicode', pretty_print=True).strip()
                parts.append(child_str)
            result_xml = '\n\n'.join(parts)
        else:
            result_xml = etree.tostring(root, encoding='unicode', pretty_print=True)
        
        if fixes_count > 0:
            print(f"[Sanitize] Удалено {fixes_count} маркеров списка из <reference>")
        
        return result_xml.strip()
        
    except Exception as e:
        print(f"[Sanitize] Ошибка санитизации: {e}")
        return xml_str

def processing_worker():
    logger.info("[Worker] Запущен поток")

    if not check_connection():
        logger.warning("[Worker] deepseek недоступен на старте")

    while True:
        task = processing_queue.get()
        if task is None:
            break

        article_id, article_content, template_id = task

        try:
            is_auto_mode = (template_id == 'auto' or template_id is None)
            template = None
            
            if not is_auto_mode:
                template = load_template(template_id)
                if not template:
                    raise ValueError(f"Template {template_id} not found")
                print(f"[Worker] 📋 Строгий режим: шаблон '{template.get('name', 'unknown')}'")
            else:
                print(f"[Worker] 🤖 Авто-режим: подбор через RAG")

            if is_auto_mode:
                current_instr = get_base_instruction()
                active_v = 1
            else:
                history = get_prompt_history(template_id)
                active_v = history.get('active', 1)
                current_instr = next((v['instruction'] for v in history['versions'] if v['v'] == active_v), get_base_instruction())

            logger.info(f"[Worker] {article_id} — Шаг 1: Генерация (Prompt V{active_v}, mode={'AUTO' if is_auto_mode else 'STRICT'})")
            
            prompt, rag_info = generate_few_shot_prompt(
                template=template,
                article_text=article_content,
                instruction=current_instr,
                use_rag=is_auto_mode
            )
            rag_info["source"] = rag_info.get("source", "unknown")
            generated_xml = run_ai_inference(prompt)

            USE_SCHEMA_VALIDATION = True

            if USE_SCHEMA_VALIDATION and generated_xml and template and template.get("after"):
                logger.info(f"[Worker] {article_id} — Шаг 2: Валидация схемы")
                try:
                    from hypothesis_schema import pipeline_validate_and_fix
                    result_xml = pipeline_validate_and_fix(
                        generated_xml_str=generated_xml,
                        example_xml_str=template["after"],
                        article_text=article_content,
                        llm_generate_fn=generate_response
                    )
                    if not result_xml:
                        result_xml = generated_xml
                except Exception as e:
                    logger.info(f"[Worker] {article_id} — Ошибка валидации схемы: {e}, использование оригинала")
                    result_xml = generated_xml
            else:
                result_xml = generated_xml

            if article_id in articles_db:
                article_data = articles_db[article_id]
                if article_data.get("cut_text") and result_xml:
                    result_xml = inject_cut_text_into_xml(result_xml, article_data["cut_text"])
                    print(f"[Worker] {article_id} — Injected cut text ({len(article_data['cut_text'])} chars)")

            if result_xml:
                try:
                    original_len = len(result_xml)
                    result_xml = ensure_valid_xml(result_xml)
                    if len(result_xml) != original_len:
                        print(f"[Worker] {article_id} — Stage 2.5: XML validation (исправлено)")
                    else:
                        print(f"[Worker] {article_id} — Stage 2.5: XML validation (OK)")
                except Exception as e:
                    print(f"[Worker] {article_id} XML validation failed: {e}")
                    
            if not TEXTUAL_GROUNDING_AVAILABLE:
                print(f"[Worker] {article_id} — Stage 3: Textual Grounding МОДУЛЬ НЕ ДОСТУПЕН")
            elif not result_xml:
                print(f"[Worker] {article_id} — Stage 3: Textual Grounding пустой XML")
            else:
                article_data = articles_db.get(article_id, {})
                original_text = article_data.get("content", "")
                
                if not original_text:
                    print(f"[Worker] {article_id} — Stage 3: Textual Grounding нет оригинального текста")
                elif len(original_text) < 100:
                    print(f"[Worker] {article_id} — Stage 3: Textual Grounding оригинал слишком короткий ({len(original_text)} симв.)")
                else:
                    try:
                        print(f"[Worker] {article_id} — Stage 3: Textual Grounding (запуск...)")
                        print(f"[Worker] {article_id}   XML: {len(result_xml)} симв., оригинал: {len(original_text)} симв.")
                        
                        result_xml, textual_fixes = textual_verify_and_fix(result_xml, original_text)
                        
                        if textual_fixes:
                            print(f"[Worker] {article_id}   Текстовых исправлений: {len(textual_fixes)}")
                            for fix in textual_fixes:
                                print(
                                    f"[Worker] {article_id}     <{fix['tag']}> "
                                    f"({fix['classification']}, sim={fix['similarity_before']:.2f})"
                                )
                        else:
                            print(f"[Worker] {article_id} Текстовых проблем не найдено (0 исправлений)")
                    except Exception as e:
                        logger.error(f"[Worker] {article_id} Textual Grounding ERROR: {e}", exc_info=True)
            if result_xml:
                try:
                    result_xml = sanitize_final_xml(result_xml)
                except Exception as e:
                    print(f"[Worker] {article_id} Sanitize failed: {e}")
                        
            if result_xml:
                processing_status["results"][article_id] = {
                    "xml": result_xml,
                    "raw_xml": generated_xml,
                    "status": "done",
                    "error": None,
                    "rag_used": rag_info.get("used", False),
                    "rag_similarity": rag_info.get("similarity", 0.0),
                    "rag_source": rag_info.get("source", "unknown"),
                    "stages": {
                        "generated": bool(generated_xml),
                        "validated": USE_SCHEMA_VALIDATION
                    }
                }
            else:
                logger.warning(f"[Worker] ✗ {article_id}: Пустой результат генерации")
                processing_status["results"][article_id] = {
                    "xml": None,
                    "raw_xml": generated_xml,
                    "status": "error",
                    "error": "Пустой результат генерации (LLM не вернул XML)",
                    "rag_used": rag_info.get("used", False),
                    "rag_similarity": rag_info.get("similarity", 0.0),
                    "rag_source": rag_info.get("source", "unknown"),
                    "stages": {
                        "generated": bool(generated_xml),
                        "validated": USE_SCHEMA_VALIDATION
                    }
                }

            if result_xml and len(result_xml) > 100:
                should_add = True
                if rag_info.get("used") and rag_info.get("similarity", 0) > 0.9:
                    should_add = False
                    logger.info(f"[Worker] RAG: Пропуск добавления (trivial match, similarity={rag_info['similarity']:.2f})")
                
                if should_add:
                    try:
                        clean_xml_for_rag = remove_text_tag_from_xml(result_xml, TEXT_TAG_NAME)
                        
                        success, msg, ex_id = add_example(
                            before=article_content,
                            after=clean_xml_for_rag,
                            source_template=template_id
                        )
                        if success:
                            size_before = len(result_xml)
                            size_after = len(clean_xml_for_rag)
                            saved_pct = int((1 - size_after/size_before) * 100) if size_before > 0 else 0
                            logger.info(f"[Worker] RAG: Добавлен пример из {article_id} (экономия {saved_pct}%)")
                        else:
                            logger.warning(f"[Worker] RAG: {msg} (id={ex_id[:8] if ex_id else 'N/A'})")
                    except Exception as e:
                        logger.error(f"[Worker] RAG: Ошибка добавления: {e}", exc_info=True)
            
            if result_xml:
                logger.info(f"[Worker] ✓ {article_id} завершена ({len(result_xml)} символов)")
            else:
                logger.warning(f"[Worker] ✗ {article_id}: завершена с ошибкой (пустой XML)")

        except ConnectionError as e:
            error_msg = f"Ошибка подключения к LLM: {e}"
            logger.error(f"[Worker] ✗ {article_id}: {error_msg}", exc_info=True)
            processing_status["results"][article_id] = {
                "xml": None, "status": "error", "error": error_msg
            }
        except Exception as e:
            error_msg = f"{type(e).__name__}: {str(e)}"
            logger.error(f"[Worker] ✗ {article_id}: {error_msg}", exc_info=True)
            processing_status["results"][article_id] = {
                "xml": None, "status": "error", "error": error_msg
            }
        finally:
            processing_status["current"] += 1
            if processing_status["current"] >= processing_status["total"]:
                processing_status["active"] = False
                logger.info("[Worker] Все задачи выполнены")
            processing_queue.task_done()


def save_template(name, before, after):
    template_id = hashlib.sha256(f"{name}{time.time()}".encode()).hexdigest()[:8]
    filepath = os.path.join(MEMORY_FOLDER, f"{template_id}.md")
    content = f"""# {name}
## BEFORE
{before}

## AFTER
{after}
"""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    return template_id

def load_template(template_id):
    if not validate_template_id(template_id):
        logger.warning(f"Попытка обхода пути в load_template: {template_id}")
        return None
    filepath = os.path.join(MEMORY_FOLDER, f"{template_id}.md")
    if not os.path.exists(filepath):
        return None
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    try:
        name = content.split('# ')[1].split('\n')[0].strip()
        parts = content.split('## BEFORE\n')
        if len(parts) < 2: return None
        rest = parts[1].split('## AFTER\n')
        if len(rest) < 2: return None
        return {
            "id": template_id,
            "name": name,
            "before": rest[0].strip(),
            "after": rest[1].strip()
        }
    except:
        return None

def list_templates():
    templates = []
    for fname in sorted(os.listdir(MEMORY_FOLDER)):
        if fname.endswith('.md'):
            tid = fname[:-3]
            tpl = load_template(tid)
            if tpl:
                templates.append({"id": tid, "name": tpl["name"]})
    return templates


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/debug/llama-status')
def debug_llama_status():
    return jsonify(get_model_info())


@app.route('/debug/test-llama', methods=['POST'])
def debug_test_llama():
    try:
        data = request.json
        prompt = data.get('prompt', 'Скажи "Привет" на русском.')

        start = time.time()
        result = generate_response(prompt)
        elapsed = time.time() - start

        return jsonify({
            "success": True,
            "time_sec": round(elapsed, 2),
            "result_length": len(result),
            "preview": result[:200] + "..." if len(result) > 200 else result
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/upload', methods=['POST'])
def upload():
    global current_md_text, articles_db
    
    if 'file' not in request.files:
        return jsonify({"error": "Файл не найден"}), 400

    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({"error": "Поддерживаются только .docx"}), 400

    stop_words = request.form.getlist('stop_words[]')
    start_words = request.form.getlist('start_words[]')
    end_words = request.form.getlist('end_words[]')
    
    preview_mode = request.form.get('preview_only', '') == '1'
    split_pattern = request.form.get('split_pattern', '').strip()
    periodicity = int(request.form.get('periodicity', '1'))
    position = request.form.get('position', 'bottom')
    test_only = request.form.get('test_only', '') == '1'

    original_filename = secure_filename(file.filename)
    if not original_filename:
        return jsonify({"error": "Некорректное имя файла"}), 400
    
    safe_filename = f"{uuid.uuid4().hex}_{original_filename}"
    filepath = os.path.join(UPLOAD_FOLDER, safe_filename)
    file.save(filepath)
    logger.info(f"Загружен файл: {safe_filename}")

    try:
        current_md_text = docx_to_markdown(filepath)

        if not current_md_text or len(current_md_text.strip()) < 10:
            return jsonify({
                "error": "Документ пуст или не содержит текста. Проверьте файл."
            }), 400

        for sw in stop_words:
            if sw and sw.upper() in current_md_text.upper():
                current_md_text = current_md_text.split(sw)[0]

        if preview_mode:
            return jsonify({
                "success": True,
                "full_text": current_md_text,
                "char_count": len(current_md_text),
                "word_count": len(current_md_text.split())
            })

        if test_only:
            articles = split_journal(current_md_text, split_pattern, periodicity, position)
            return jsonify({
                "articles": [{"id": f"art_{i}", "title": a["title"]} for i, a in enumerate(articles)],
                "count": len(articles)
            })

        articles = split_journal(current_md_text, split_pattern, periodicity, position)
        articles_db.clear()
        
        for i, art in enumerate(articles):
            aid = f"art_{i}"
            filtered = remove_dynamic_block(art["content"], start_words, end_words)

            cut_text = None
            if filtered != art["content"] and start_words and end_words:
                orig = art["content"]
                start_pos = None
                for sw in start_words:
                    idx = orig.lower().find(sw.lower())
                    if idx != -1:
                        start_pos = idx
                        break
                if start_pos is not None:
                    after_start = orig[start_pos:]
                    for ew in end_words:
                        esc = re.escape(ew)
                        pattern = f"{esc.replace(r'\\s+', r'\\\\s+')}(?![а-яА-ЯёЁ,;:])"
                        m = re.search(pattern, after_start, re.IGNORECASE)
                        if m:
                            cut_end = start_pos + m.start()
                            raw_cut = orig[start_pos:cut_end].strip()
                            cut_text = re.sub(r'\s+', ' ', raw_cut).strip()
                            break

            articles_db[aid] = {
                "title": art["title"],
                "content": art["content"],
                "filtered": filtered,
                "cut_text": cut_text,
                "parsed": None
            }

        article_ids = [{"id": aid, "title": articles_db[aid]["title"]} for aid in articles_db]
        logger.info(f"[Upload] Загружен журнал: {len(articles)} статей, размер: {len(current_md_text)} симв.")
        return jsonify({"articles": article_ids})

    except Exception as e:
        logger.error(f"[Upload] Ошибка при загрузке файла: {e}", exc_info=True)
        return jsonify({"error": "Внутренняя ошибка сервера при обработке документа."}), 500
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)

@app.route('/prompt/base', methods=['GET', 'POST'])
def base_instruction_endpoint():
    if request.method == 'POST':
        data = request.json
        instruction = data.get('instruction', '').strip()
        if not instruction:
            return jsonify({"error": "Промпт не может быть пустым"}), 400
        save_base_instruction(instruction)
        return jsonify({"success": True})
    return jsonify({"instruction": get_base_instruction()})


@app.route('/prompt/base/reset', methods=['POST'])
def reset_base_instruction():
    filepath = os.path.join(MEMORY_FOLDER, "base_instruction.json")
    if os.path.exists(filepath):
        os.remove(filepath)
    return jsonify({"success": True, "instruction": DEFAULT_INSTRUCTION})

@app.route('/article/<aid>')
def get_article(aid):
    if aid in articles_db:
        return jsonify(articles_db[aid])
    return jsonify({"error": "Статья не найдена"}), 404


@app.route('/memory', methods=['GET', 'POST'])
def memory_endpoint():
    if request.method == 'POST':
        data = request.json
        name = data.get('name', 'Без названия')
        before = data.get('before', '')
        after = data.get('after', '')
        if not before.strip() or not after.strip():
            return jsonify({"error": "Оба поля должны быть заполнены"}), 400
        tid = save_template(name, before, after)
    
        try:
            from rag_engine import add_example, remove_text_tag_from_xml
            clean_after = remove_text_tag_from_xml(after, TEXT_TAG_NAME)
            success, msg, ex_id = add_example(
                before=before,
                after=clean_after,
                source_template=tid
            )
            bootstrap_info = {"success": success, "message": msg, "example_id": ex_id}
            logger.info(f"[Bootstrap] Шаблон '{name}' добавлен в RAG-базу: {msg}")
        except Exception as e:
            bootstrap_info = {"success": False, "message": str(e)}
            logger.error(f"[Bootstrap] Ошибка добавления в RAG: {e}", exc_info=True)
        
        return jsonify({
            "id": tid, 
            "name": name,
            "bootstrap": bootstrap_info
        })
    return jsonify(list_templates())


@app.route('/memory/<template_id>', methods=['GET', 'DELETE'])
def template_endpoint(template_id):
    if not validate_template_id(template_id):
        return jsonify({"error": "Некорректный ID шаблона"}), 400
    if request.method == 'DELETE':
        filepath = os.path.join(MEMORY_FOLDER, f"{template_id}.md")
        if os.path.exists(filepath):
            os.remove(filepath)
            return jsonify({"success": True})
        return jsonify({"error": "Шаблон не найден"}), 404
    tpl = load_template(template_id)
    if tpl:
        return jsonify(tpl)
    return jsonify({"error": "Шаблон не найден"}), 404


@app.route('/parse/start', methods=['POST'])
def start_parsing():
    global processing_status
    data = request.json
    template_id = data.get('template_id')
    article_ids = data.get('article_ids', [])

    logger.info(f"[Parse] Запуск обработки: template_id={template_id}, articles={len(article_ids)}")
    
    valid_ids = []
    skipped = {"hidden": 0, "deleted": 0}
    
    for aid in article_ids:
        if aid not in articles_db:
            continue
        art = articles_db[aid]
        if art.get("deleted"):
            skipped["deleted"] += 1
            continue
        if art.get("hidden"):
            skipped["hidden"] += 1
            continue
        valid_ids.append(aid)
    
    if skipped["hidden"] > 0 or skipped["deleted"] > 0:
        logger.info(f"[Parse] Пропущено: скрыто={skipped['hidden']}, удалено={skipped['deleted']}")
    
    if not valid_ids:
        return jsonify({"error": "Нет статей для обработки"}), 400
    
    if not template_id:
        return jsonify({"error": "Не выбран режим обработки (шаблон или Авто)"}), 400

    processing_status = {
        "active": True,
        "current": 0,
        "total": len(valid_ids),
        "results": {},
        "template_id": template_id
    }

    for aid in valid_ids:
        content = articles_db[aid].get("edited_content") or articles_db[aid]["filtered"]
        processing_queue.put((aid, content, template_id))

    return jsonify({
        "status": "started", 
        "total": len(valid_ids),
        "skipped": skipped
    })


@app.route('/parse/status')
def parsing_status():
    return jsonify(processing_status)


@app.route('/parse/result/<aid>')
def get_parsed_result(aid):
    if aid in articles_db and articles_db[aid].get("parsed"):
        return jsonify({
            "xml": articles_db[aid]["parsed"],
            "rag_used": False,
            "rag_similarity": 0.0,
            "rag_source": "cached"
        })
    if aid in processing_status["results"]:
        res = processing_status["results"][aid]
        if res["status"] == "done" and res["xml"]:
            articles_db[aid]["parsed"] = res["xml"]
            return jsonify({
                "xml": res["xml"],
                "rag_used": res.get("rag_used", False),
                "rag_similarity": res.get("rag_similarity", 0.0),
                "rag_source": res.get("rag_source", "unknown")
            })
        elif res["status"] == "error":
            return jsonify({"error": res["error"]}), 500
    return jsonify({"status": "pending"}), 202


@app.route('/parse/all-results')
def get_all_parsed():
    results = []
    for aid in sorted(articles_db.keys()):
        parsed = articles_db[aid].get("parsed")
        if parsed:
            validated = ensure_valid_xml(parsed)
            sanitized = sanitize_final_xml(validated)
            results.append(sanitized)
    
    if not results:
        return jsonify({"error": "Нет готовых результатов"}), 404
    
    combined = "\n\n".join(results)
    combined = re.sub(r'\n{3,}', '\n\n', combined)
    return jsonify({"combined_xml": combined})


@app.route('/prompt/history/<template_id>')
def get_prompt_history_endpoint(template_id):
    return jsonify(get_prompt_history(template_id))

def handle_auto_mode_correction(article_id, remark, current_xml, original_text=""):
    if not article_id or article_id not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    if not current_xml:
        return jsonify({"error": "Нет XML для коррекции"}), 400
    
    pattern = rf'<{TEXT_TAG_NAME}[^>]*>(.*?)</{TEXT_TAG_NAME}>'
    match = re.search(pattern, current_xml, flags=re.DOTALL)
    extracted_text = None
    slim_xml = current_xml
    
    if match:
        extracted_text = match.group(1)
        slim_xml = re.sub(pattern, '', current_xml, flags=re.DOTALL)
        slim_xml = re.sub(r'\n\s*\n\s*\n+', '\n\n', slim_xml).strip()
        print(f"[AutoRefine] Вырезан <{TEXT_TAG_NAME}> из XML перед отправкой в LLM ({len(extracted_text)} симв.)")
    
    correction_prompt = f"""Ты — эксперт по JATS-XML разметке научных статей.

ТЕКУЩИЙ XML-РЕЗУЛЬТАТ (тег <{TEXT_TAG_NAME}> с текстом статьи удалён для экономии контекста):
{slim_xml}

ЗАМЕЧАНИЕ ПОЛЬЗОВАТЕЛЯ:
"{remark}"

ЗАДАЧА:
Внеси правки в XML-метаданные в соответствии с замечанием. Верни ТОЛЬКО исправленный XML-документ.
НЕ добавляй тег <{TEXT_TAG_NAME}> обратно!
Не добавляй комментариев, не оборачивай в кодовые блоки (без ```xml)."""
    
    try:
        corrected_slim_xml = run_ai_inference(correction_prompt)
        corrected_slim_xml = corrected_slim_xml.strip().strip('`')
        corrected_slim_xml = corrected_slim_xml.replace('```xml', '').replace('```', '').strip()
        
        if not corrected_slim_xml or len(corrected_slim_xml) < 50:
            return jsonify({"error": "LLM вернула пустой результат"}), 500
        
        if extracted_text:
            corrected_xml = inject_cut_text_into_xml(corrected_slim_xml, extracted_text)
            print(f"[AutoRefine] Тег <{TEXT_TAG_NAME}> возвращён на место после коррекции")
        else:
            corrected_xml = corrected_slim_xml
        
        article_data = articles_db[article_id]
        full_original = article_data.get("content", "")
        textual_fixes = []
        
        if TEXTUAL_GROUNDING_AVAILABLE and full_original and len(full_original) > 100:
            try:
                print(f"[AutoRefine] {article_id} — Textual Grounding (запуск...)")
                print(f"[AutoRefine] {article_id}   XML: {len(corrected_xml)} симв., оригинал: {len(full_original)} симв.")
                
                corrected_xml, textual_fixes = textual_verify_and_fix(corrected_xml, full_original)
                
                if textual_fixes:
                    print(f"[AutoRefine] {article_id}   ✅ Текстовых исправлений: {len(textual_fixes)}")
                    for fix in textual_fixes[:5]:
                        print(
                            f"[AutoRefine] {article_id}     ✓ <{fix['tag']}> "
                            f"({fix['classification']}, sim={fix['similarity_before']:.2f})"
                        )
                    if len(textual_fixes) > 5:
                        print(f"[AutoRefine] {article_id}     ... и ещё {len(textual_fixes) - 5} исправлений")
                else:
                    print(f"[AutoRefine] {article_id}   ✓ Текстовых проблем не найдено")
            except Exception as e:
                import traceback
                print(f"[AutoRefine] {article_id} ❌ Textual Grounding ERROR: {e}")
                traceback.print_exc()
        
        articles_db[article_id]["parsed"] = corrected_xml
        if article_id in processing_status["results"]:
            processing_status["results"][article_id]["xml"] = corrected_xml
        
        update_rag_example_for_article(article_id, corrected_xml)
        
        return jsonify({
            "success": True,
            "new_version": "corrected",
            "msg": f"XML исправлен. Валидация наполнения тэгов: {len(textual_fixes)} исправлений.",
            "textual_fixes_count": len(textual_fixes)
        })
        
    except Exception as e:
        logger.error(f"[AutoRefine] Ошибка: {e}", exc_info=True)
        return jsonify({"error": "Внутренняя ошибка при коррекции XML."}), 500


def update_rag_example_for_article(article_id, new_xml):
    if article_id not in articles_db:
        return
    
    article_data = articles_db[article_id]
    before = article_data.get("edited_content") or article_data.get("filtered") or article_data.get("content", "")
    
    if not before:
        return
    
    clean_xml = remove_text_tag_from_xml(new_xml, TEXT_TAG_NAME)
    
    example_id = hashlib.sha256(before[:3000].encode('utf-8')).hexdigest()[:16]
    
    db = get_rag_db()
    found = False
    
    for ex in db["examples"]:
        if ex["id"] == example_id:
            ex["after"] = clean_xml
            ex["timestamp"] = time.strftime("%Y-%m-%d %H:%M:%S")
            ex["version"] = ex.get("version", 1) + 1
            ex["source_template"] = "auto_learned"
            found = True
            print(f"[RAG] 🔄 Обновлён пример {example_id[:8]}... (v{ex['version']})")
            break
    
    if not found:
        db["examples"].append({
            "id": example_id,
            "before": before,
            "after": clean_xml,
            "source_template": "auto_learned",
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "version": 1
        })
        print(f"[RAG] ➕ Добавлен auto-learned пример {example_id[:8]}...")
    
    save_rag_db(db)

@app.route('/prompt/refine', methods=['POST'])
def refine_prompt():
    data = request.json
    template_id = data.get('template_id')
    article_id = data.get('article_id')
    remark = data.get('remark')
    apply_scope = data.get('apply_scope', 'article')
    current_xml = data.get('current_xml', '')
    
    if template_id == 'auto' or not template_id:
        original_text = data.get('original_text', '')
        return handle_auto_mode_correction(article_id, remark, current_xml, original_text)
    
    if not template_id:
        return jsonify({"error": "Не выбран шаблон"}), 400
        
    history = get_prompt_history(template_id)
    active_v = history.get('active', 1)
    current_instr = next((v['instruction'] for v in history['versions'] if v['v'] == active_v), DEFAULT_INSTRUCTION)
    
    meta_prompt = f"""Ты — эксперт по промпт-инжинирингу и лингвистике. Твоя задача — улучшить системную инструкцию для ИИ-библиографа.

ТЕКУЩАЯ ИНСТРУКЦИЯ:
---
{current_instr}
---

ЗАМЕЧАНИЕ ПОЛЬЗОВАТЕЛЯ:
"{remark}"

ТВОЯ ЗАДАЧА:
Перепиши ТЕКУЩУЮ ИНСТРУКЦИЮ, добавив или уточнив правила, чтобы учесть замечание пользователя.
ВАЖНО: Верни ТОЛЬКО обновленный текст инструкции. БЕЗ префиксов вроде "Вот обновленная инструкция", БЕЗ комментариев, БЕЗ оформления в markdown. Просто чистый текст."""

    try:
        new_instr = run_ai_inference(meta_prompt)
        new_instr = new_instr.strip().strip('`').replace('markdown', '').replace('text', '').strip()
        if new_instr.lower().startswith("вот обновленная"):
            new_instr = new_instr.split("\n", 1)[-1].strip()
            
        new_v = max(v['v'] for v in history['versions']) + 1
        history['versions'].append({
            "v": new_v,
            "instruction": new_instr,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "remark": remark
        })
        history['active'] = new_v
        save_prompt_history(template_id, history)
        
        if apply_scope == 'article' and article_id:
            restart_parsing([article_id], template_id)
            print(f"[Refine] Перезапущена статья {article_id} с новой версией V{new_v}")
            return jsonify({"success": True, "new_version": new_v, "msg": f"Перезапущена текущая статья (V{new_v})"})
        elif apply_scope == 'journal':
            restart_parsing(list(articles_db.keys()), template_id)
            print(f"[Refine] Перезапущен весь журнал с новой версией V{new_v}")
            return jsonify({"success": True, "new_version": new_v, "msg": f"Перезапущен весь журнал (V{new_v})"})
        else:
            print(f"[Refine] Промпт V{new_v} сохранён для будущих запусков")
            return jsonify({"success": True, "new_version": new_v, "msg": "Промпт сохранен для будущих запусков"})
    except Exception as e:
        logger.error(f"[Refine] Ошибка: {e}", exc_info=True)
        return jsonify({"error": "Внутренняя ошибка при обновлении промпта."}), 500

@app.route('/prompt/rollback', methods=['POST'])
def rollback_prompt():
    data = request.json
    template_id = data.get('template_id')
    target_v = data.get('version')
    
    history = get_prompt_history(template_id)
    if any(v['v'] == target_v for v in history['versions']):
        history['active'] = target_v
        save_prompt_history(template_id, history)
        return jsonify({"success": True, "active": target_v})
    return jsonify({"error": "Version not found"}), 404

from difflib import SequenceMatcher

def normalize_text(text):
    text = text.replace('\u00A0', ' ')
    text = text.replace('–', '-').replace('—', '-').replace('−', '-').replace('―', '-')
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def tokenize(text):
    return re.findall(r'\w+|[^\w\s]', text, re.UNICODE)

def is_word_token(token):
    if token.isdigit():
        return False
    if len(token) == 1 and not token.isalnum():
        return False
    return any(c.isalpha() for c in token) and len(token) > 1

def find_common_phrases(tokenized_examples, min_tokens=3):
    base = tokenized_examples[0]
    constants = []
    
    i = 0
    while i < len(base):
        best_seq = []
        max_look = min(i + 40, len(base) + 1)
        
        for j in range(i + min_tokens, max_look):
            candidate = base[i:j]
            if len(candidate) < min_tokens:
                continue
            
            found_in_all = True
            for other in tokenized_examples[1:]:
                found = any(
                    other[m:m+len(candidate)] == candidate
                    for m in range(len(other) - len(candidate) + 1)
                )
                if not found:
                    found_in_all = False
                    break
            
            if found_in_all:
                best_seq = candidate
            else:
                break
        
        if len(best_seq) >= min_tokens:
            constants.append(best_seq)
            i += len(best_seq)
        else:
            i += 1
    
    return constants

def build_regex(constants, tokenized_first):
    parts = []
    
    for idx, const in enumerate(constants):
        if idx == 0:
            if tokenized_first[:len(const)] != const:
                parts.append('.*?')
        else:
            parts.append('.*?')
        
        const_parts = []
        for i, token in enumerate(const):
            if i > 0:
                prev = const[i-1]
                if is_word_token(prev) and is_word_token(token):
                    const_parts.append(r'\s+')
                else:
                    const_parts.append(r'\s*')
            const_parts.append(re.escape(token))
        
        parts.append(''.join(const_parts))
    
    if tokenized_first[-len(constants[-1]):] != constants[-1]:
        parts.append('.*?')
    
    return ''.join(parts)

def synthesize_regex_from_examples(examples):
    if len(examples) < 2:
        return None, "Нужно минимум 2 примера"
    
    normalized = [normalize_text(ex) for ex in examples]
    tokenized = [tokenize(n) for n in normalized]
    
    constants = find_common_phrases(tokenized, min_tokens=3)
    
    if not constants:
        return None, "Не найдено общих фраз (≥ 3 токенов). Убедитесь, что примеры однотипны."
    
    regex_pattern = build_regex(constants, tokenized[0])
    
    try:
        compiled = re.compile(regex_pattern, re.DOTALL | re.IGNORECASE)
        for i, ex in enumerate(examples):
            if not compiled.search(ex) and not compiled.search(normalized[i]):
                return None, f"Паттерн не нашёл пример #{i+1}"
    except re.error as e:
        return None, f"Ошибка компиляции regex: {e}"
    
    const_phrases = [' '.join(c) for c in constants[:5]]
    explanation = (
        f"Паттерн синтезирован из {len(examples)} примеров. "
        f"Найдено {len(constants)} общих фраз: " +
        ", ".join(f'«{p}»' for p in const_phrases) +
        (f" и ещё {len(constants) - 5}" if len(constants) > 5 else "") +
        "."
    )
    
    return regex_pattern, explanation


@app.route('/synthesize-pattern', methods=['POST'])
def synthesize_pattern():
    data = request.json
    examples = data.get('examples', [])
    
    if len(examples) < 2:
        return jsonify({"error": "Нужно минимум 2 примера"}), 400
    
    if len(examples) > 5:
        return jsonify({"error": "Максимум 5 примеров"}), 400
    
    try:
        regex_pattern, explanation = synthesize_regex_from_examples(examples)
        
        if not regex_pattern:
            return jsonify({"error": explanation}), 400
        
        match_count = None
        if current_md_text:
            try:
                matches = list(re.finditer(regex_pattern, current_md_text, re.DOTALL | re.IGNORECASE))
                match_count = len(matches)
            except Exception:
                pass
        
        return jsonify({
            "success": True,
            "regex": regex_pattern,
            "explanation": explanation,
            "matches_found": match_count,
            "examples_used": len(examples)
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Ошибка синтеза: {str(e)}"}), 500

@app.route('/full-document')
def get_full_document():
    if not current_md_text:
        return jsonify({"error": "Документ не загружен"}), 404
    return jsonify({
        "text": current_md_text,
        "char_count": len(current_md_text),
        "word_count": len(current_md_text.split())
    })

@app.route('/rag/stats')
def rag_stats():
    return jsonify(get_db_stats())

@app.route('/rag/add', methods=['POST'])
def rag_add_example():
    data = request.json
    before = data.get('before', '').strip()
    after = data.get('after', '').strip()
    source = data.get('source_template')
    
    if not before or not after:
        return jsonify({"error": "Нужны before и after"}), 400
    
    success, msg, ex_id = add_example(before, after, source)
    
    return jsonify({
        "success": success,
        "message": msg,
        "example_id": ex_id
    })


@app.route('/rag/examples')
def rag_list_examples():
    db = get_rag_db()
    result = []
    for ex in db["examples"]:
        result.append({
            "id": ex["id"],
            "before_preview": ex["before"][:200] + "..." if len(ex["before"]) > 200 else ex["before"],
            "after_preview": ex["after"][:200] + "..." if len(ex["after"]) > 200 else ex["after"],
            "source_template": ex.get("source_template"),
            "timestamp": ex.get("timestamp"),
            "before_length": len(ex["before"]),
            "after_length": len(ex["after"])
        })
    return jsonify(result)


@app.route('/rag/example/<example_id>')
def rag_get_example(example_id):
    db = get_rag_db()
    for ex in db["examples"]:
        if ex["id"] == example_id:
            return jsonify(ex)
    return jsonify({"error": "Пример не найден"}), 404


@app.route('/rag/example/<example_id>', methods=['DELETE'])
def rag_delete_example(example_id):
    from rag_engine import remove_example
    if remove_example(example_id):
        return jsonify({"success": True})
    return jsonify({"error": "Пример не найден"}), 404


@app.route('/rag/test-similarity', methods=['POST'])
def rag_test_similarity():
    data = request.json
    text = data.get('text', '')
    threshold = float(data.get('threshold', 0.25))
    
    if not text:
        return jsonify({"error": "Пустой текст"}), 400
    
    match = find_similar_example(text, threshold=threshold)
    return jsonify({
        "found": match is not None,
        "match": match
    })

@app.route('/settings/text-tag', methods=['GET', 'POST'])
def text_tag_settings():
    global TEXT_TAG_NAME
    
    if request.method == 'POST':
        data = request.json
        new_tag = data.get('tag_name', '').strip()
        if not new_tag:
            return jsonify({"error": "Имя тега не может быть пустым"}), 400
        if not re.match(r'^[a-zA-Z][a-zA-Z0-9_-]*$', new_tag):
            return jsonify({"error": "Некорректное имя тега. Используйте только латинские буквы, цифры, _ и -"}), 400
        
        TEXT_TAG_NAME = new_tag
        print(f"[Settings] Имя тега основного текста изменено на: <{new_tag}>")
        return jsonify({"success": True, "tag_name": new_tag})
    
    return jsonify({"tag_name": TEXT_TAG_NAME})

@app.route('/article/edit/<aid>', methods=['POST'])
def edit_article(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    data = request.json
    new_content = data.get('content', '').strip()
    
    if not new_content:
        return jsonify({"error": "Содержимое не может быть пустым"}), 400
    
    articles_db[aid]["edited_content"] = new_content
    
    articles_db[aid]["parsed"] = None
    if aid in processing_status["results"]:
        del processing_status["results"][aid]
    
    logger.info(f"[Curation] Статья {aid} отредактирована ({len(new_content)} симв.)")
    
    return jsonify({
        "success": True,
        "content_length": len(new_content)
    })

@app.route('/article/edit-xml/<aid>', methods=['POST'])
def edit_article_xml(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    data = request.json
    new_xml = (data.get('xml') or '').strip()
    
    if not new_xml:
        return jsonify({"error": "XML не может быть пустым"}), 400
    
    xml_parseable = True
    try:
        from lxml import etree
        wrapped = f"<_root_>{new_xml}</_root_>"
        secure_parser = etree.XMLParser(resolve_entities=False, no_network=True)
        etree.fromstring(wrapped.encode('utf-8'), parser=secure_parser)
    except Exception as e1:
        try:
            parser = etree.XMLParser(recover=True, resolve_entities=False, no_network=True)
            etree.fromstring(wrapped.encode('utf-8'), parser)
        except Exception:
            xml_parseable = False

    articles_db[aid]["parsed"] = new_xml
    articles_db[aid]["xml_manually_edited"] = True
    
    if aid in processing_status["results"]:
        processing_status["results"][aid]["xml"] = new_xml
        processing_status["results"][aid]["manually_edited"] = True
        processing_status["results"][aid]["xml_valid"] = xml_parseable
    
    validity_note = " (валидный XML)" if xml_parseable else " (возможны синтаксические ошибки)"
    print(f"[ManualEdit] ✏️ {aid}: XML обновлён вручную ({len(new_xml)} симв.){validity_note}")
    
    return jsonify({
        "success": True,
        "length": len(new_xml),
        "xml_valid": xml_parseable
    })

@app.route('/article/hide/<aid>', methods=['POST'])
def hide_article(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    articles_db[aid]["hidden"] = True
    logger.info(f"[Curation] Статья {aid} скрыта")
    return jsonify({"success": True})


@app.route('/article/restore/<aid>', methods=['POST'])
def restore_article(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    articles_db[aid]["hidden"] = False
    print(f"[Curation] Статья {aid} восстановлена")
    return jsonify({"success": True})


@app.route('/article/delete/<aid>', methods=['POST'])
def delete_article(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    articles_db[aid]["deleted"] = True
    logger.info(f"[Curation] Статья {aid} удалена")
    return jsonify({"success": True})


@app.route('/article/undelete/<aid>', methods=['POST'])
def undelete_article(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    articles_db[aid]["deleted"] = False
    print(f"[Curation] Статья {aid} восстановлена из удалённых")
    return jsonify({"success": True})


@app.route('/article/reset-edits/<aid>', methods=['POST'])
def reset_article_edits(aid):
    if aid not in articles_db:
        return jsonify({"error": "Статья не найдена"}), 404
    
    if "edited_content" in articles_db[aid]:
        del articles_db[aid]["edited_content"]
        articles_db[aid]["parsed"] = None
        if aid in processing_status["results"]:
            del processing_status["results"][aid]
        print(f"[Curation] Статья {aid} сброшена к оригиналу")
    
    return jsonify({"success": True})


@app.route('/articles/state')
def get_articles_state():
    state = {}
    for aid, data in articles_db.items():
        state[aid] = {
            "hidden": data.get("hidden", False),
            "deleted": data.get("deleted", False),
            "edited": "edited_content" in data,
            "parsed": data.get("parsed") is not None
        }
    return jsonify(state)

if __name__ == '__main__':
    logger.info(f"[Main] Подключение к LLM: http://localhost:8080")
    
    if TEXTUAL_GROUNDING_AVAILABLE:
        logger.info(f"[Main] textual_grounding.py загружен успешно")
    else:
        logger.error(f"[Main] textual_grounding.py НЕ ЗАГРУЖЕН")
    
    threading.Thread(target=processing_worker, daemon=True).start()
    logger.info("[Main] Запущен рабочий поток")

    is_production = os.environ.get("FLASK_ENV") == "production"
    app.run(debug=not is_production, threaded=True, use_reloader=False)
